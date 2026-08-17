#!/usr/bin/env python3
"""
Tạo file Word (.docx) cho báo cáo HUKI EBOOK
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set background color for table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_number(doc, text, level):
    """Add heading with number"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def create_table_with_header(doc, headers, rows, header_color="1E4D78"):
    """Create table with colored header"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(header_cells[i], header_color)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            row_cells[col_idx].text = str(cell_data)

    return table

def main():
    doc = Document()

    # ==================== PAGE SETUP ====================
    sections = doc.sections
    for section in sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # ==================== COVER PAGE ====================
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("HỆ THỐNG THƯƠNG MẠI ĐIỆN TỬ\nHỖ TRỢ BÁN SÁCH VẬT LÝ VÀ SÁCH ĐIỆN TỬ")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(30, 77, 120)

    doc.add_paragraph()
    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("HUKI EBOOK")
    sub_run.bold = True
    sub_run.font.size = Pt(28)
    sub_run.font.color.rgb = RGBColor(0, 123, 182)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Author info
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run("ĐỀ TÀI CUỐI KHÓA\n\nSinh viên thực hiện:\nHuỳnh Gia Huy\nLê Đức Kiên\n\nNgày: 17/08/2026")
    info_run.font.size = Pt(14)

    # Page break
    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    doc.add_heading("MỤC LỤC", level=1)

    toc_items = [
        ("1.", "GIỚI THIỆU", "3"),
        ("1.1", "Giới thiệu đề tài", "3"),
        ("1.2", "Mục tiêu đề tài", "3"),
        ("1.3", "Phạm vi dự án", "4"),
        ("2.", "KIẾN TRÚC HỆ THỐNG", "5"),
        ("2.1", "Kiến trúc Microservices", "5"),
        ("2.2", "Chi tiết các Services", "6"),
        ("3.", "CÔNG NGHỆ SỬ DỤNG", "7"),
        ("3.1", "Backend Stack", "7"),
        ("3.2", "Frontend Stack", "7"),
        ("3.3", "Cloud & Infrastructure", "8"),
        ("3.4", "Thanh toán", "8"),
        ("4.", "THIẾT KẾ DATABASE", "9"),
        ("4.1", "PostgreSQL Databases", "9"),
        ("4.2", "MongoDB Database", "10"),
        ("5.", "DEPLOYMENT", "11"),
        ("5.1", "Kiến trúc Infrastructure trên AWS", "11"),
        ("5.2", "External Cloud Services", "12"),
        ("5.3", "Security Architecture", "12"),
        ("6.", "KẾT LUẬN", "13"),
        ("6.1", "Tổng kết", "13"),
        ("6.2", "Hướng phát triển tương lai", "13"),
    ]

    for num, title, page in toc_items:
        p = doc.add_paragraph()
        p.add_run(f"{num}\t{title}").font.size = Pt(12)
        p.add_run(f"\t{'.' * 50}\t{page}").font.size = Pt(12)

    doc.add_page_break()

    # ==================== SECTION 1: INTRODUCTION ====================
    doc.add_heading("1. GIỚI THIỆU", level=1)

    doc.add_heading("1.1. Giới thiệu đề tài", level=2)

    intro_text = """Trong bối cảnh thương mại điện tử phát triển mạnh mẽ tại Việt Nam, ngành công nghiệp sách đang chuyển đổi số với xu hướng đọc sách trực tuyến ngày càng tăng. Tuy nhiên, hiện tại chưa có nền tảng TMĐT nào tại Việt Nam hỗ trợ đồng thời cả việc bán sách vật lý (có tồn kho, giao hàng) lẫn sách điện tử (PDF/EPUB) một cách toàn diện."""
    doc.add_paragraph(intro_text)

    doc.add_paragraph("Vấn đề đặt ra:", style='List Bullet')
    doc.add_paragraph("Người đọc phải sử dụng nhiều nền tảng khác nhau để mua sách vật lý và sách điện tử", style='List Bullet')
    doc.add_paragraph("Thiếu hệ thống Marketplace cho phép nhiều cửa hàng bán sách trên cùng một nền tảng", style='List Bullet')
    doc.add_paragraph("Chưa có giải pháp tích hợp đọc sách điện tử trực tiếp trên nền tảng mua sắm", style='List Bullet')

    doc.add_paragraph("Giải pháp đề xuất: HUKI EBOOK - Hệ thống TMĐT tích hợp, cho phép người bán đăng bán cả sách vật lý lẫn sách điện tử, người mua có thể mua sắm, thảo luận và đọc sách trực tiếp trên một nền tảng duy nhất.")

    doc.add_heading("1.2. Mục tiêu đề tài", level=2)

    doc.add_paragraph("Mục tiêu chính:", style='List Bullet')
    doc.add_paragraph("Xây dựng hệ thống thương mại điện tử hỗ trợ bán sách vật lý và sách điện tử hoàn chỉnh.", style='List Bullet')

    doc.add_paragraph("Mục tiêu cụ thể:", style='List Bullet')

    objectives = [
        "Hệ thống Marketplace cho phép nhiều cửa hàng đăng bán sách trên nền tảng",
        "Hỗ trợ bán sách vật lý với quản lý tồn kho và giao hàng",
        "Hỗ trợ bán sách điện tử (PDF, EPUB) với tính năng đọc trực tiếp",
        "Diễn đàn thảo luận (Forum) giữa người đọc",
        "Chat real-time giữa người mua và người bán",
        "Thanh toán trực tuyến qua PayOS và COD"
    ]

    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')

    doc.add_heading("1.3. Phạm vi dự án", level=2)

    # Create table for scope
    scope_data = [
        ["Trong phạm vi", "Ngoài phạm vi"],
        ["Backend: NestJS Microservices", "Mobile Application"],
        ["Frontend: Next.js 14 Web Application", "Triển khai Production"],
        ["Database: PostgreSQL, MongoDB", "Tính năng AI/ML"],
        ["Cloud Infrastructure: AWS EC2, Cloudflare", ""]
    ]

    table = doc.add_table(rows=len(scope_data), cols=2)
    table.style = 'Table Grid'

    for i, row_data in enumerate(scope_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            if i == 0:
                row.cells[j].paragraphs[0].runs[0].bold = True
                set_cell_shading(row.cells[j], "1E4D78")
                row.cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    doc.add_page_break()

    # ==================== SECTION 2: ARCHITECTURE ====================
    doc.add_heading("2. KIẾN TRÚC HỆ THỐNG", level=1)

    doc.add_heading("2.1. Kiến trúc Microservices", level=2)

    arch_text = """Hệ thống được xây dựng theo kiến trúc Microservices với 6 services chính, mỗi service đảm nhận một nhóm chức năng riêng biệt, giúp dễ dàng mở rộng và bảo trì."""
    doc.add_paragraph(arch_text)

    doc.add_paragraph("Sơ đồ kiến trúc:")

    # Architecture diagram as text
    arch_diagram = """
┌─────────────────────────────────────────────────────────────────┐
│                         CLIENTS                                   │
│                    Web Browser (Next.js)                          │
└─────────────────────────────┬───────────────────────────────────┘
                              │
                              ▼
┌─────────────────────────────────────────────────────────────────┐
│                       API GATEWAY (:3000)                        │
│              Auth │ Routing │ Rate Limit │ CORS                 │
└─────────────────────────────┬───────────────────────────────────┘
                              │
      ┌───────────┬───────────┼───────────┬───────────┬───────────┐
      │           │           │           │           │           │
      ▼           ▼           ▼           ▼           ▼           ▼
┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐ ┌──────────┐
│ Identity │ │Business  │ │Commerce  │ │Shipping  │ │Communi-  │ │Promo-   │
│ Service  │ │ Service  │ │ Service  │ │ Service  │ │ty        │ │tion     │
│  :3001   │ │  :3002   │ │  :3003   │ │  :3004   │ │Service   │ │Service  │
│          │ │          │ │          │ │          │ │  :3005   │ │  :3007  │
└────┬─────┘ └────┬─────┘ └────┬─────┘ └────┬─────┘ └────┬─────┘ └────┬─────┘
     │            │            │            │            │            │
     ▼            ▼            ▼            ▼            ▼            ▼
┌────────┐  ┌────────┐  ┌────────┐  ┌────────┐  ┌────────┐  ┌────────┐
│Postgre │  │Postgre │  │Postgre │  │Postgre │  │ Mongo  │  │Postgre │
│SQL     │  │SQL     │  │SQL     │  │SQL     │  │DB      │  │SQL     │
└────────┘  └────────┘  └────────┘  └────────┘  └────────┘  └────────┘
"""
    diagram_para = doc.add_paragraph()
    diagram_run = diagram_para.add_run(arch_diagram)
    diagram_run.font.name = 'Courier New'
    diagram_run.font.size = Pt(8)

    doc.add_heading("2.2. Chi tiết các Services", level=2)

    services_headers = ["Service", "Port", "Database", "Mô tả chức năng"]
    services_data = [
        ["API Gateway", "3000", "-", "Điều hướng request, xác thực JWT, rate limiting, CORS"],
        ["Identity Service", "3001", "PostgreSQL", "Xác thực người dùng, quản lý tài khoản, JWT tokens"],
        ["Business Service", "3002", "PostgreSQL", "Quản lý doanh nghiệp, cửa hàng, thành viên"],
        ["Commerce Service", "3003", "PostgreSQL", "Quản lý sách, giỏ hàng, đơn hàng, thanh toán"],
        ["Shipping Service", "3004", "PostgreSQL", "Quản lý vận chuyển, giao hàng"],
        ["Community Service", "3005", "MongoDB", "Diễn đàn, chat real-time, đánh giá, thông báo"],
        ["Promotion Service", "3007", "PostgreSQL", "Voucher, khuyến mãi, giảm giá"]
    ]

    create_table_with_header(doc, services_headers, services_data)

    doc.add_page_break()

    # ==================== SECTION 3: TECHNOLOGY ====================
    doc.add_heading("3. CÔNG NGHỆ SỬ DỤNG", level=1)

    doc.add_heading("3.1. Backend Stack", level=2)

    backend_headers = ["Component", "Technology", "Purpose"]
    backend_data = [
        ["Framework", "NestJS", "Xây dựng Microservices backend"],
        ["Language", "TypeScript", "Đảm bảo type safety"],
        ["Database ORM", "Prisma", "Truy cập database, migrations"],
        ["Primary Database", "PostgreSQL", "Lưu trữ dữ liệu chính"],
        ["Document Database", "MongoDB", "Lưu trữ forum, chat, reviews"],
        ["Cache", "Redis", "Session caching, rate limiting"],
        ["Message Queue", "RabbitMQ", "Xử lý sự kiện bất đồng bộ"],
        ["Real-time", "Socket.IO", "Chat real-time"],
        ["API Documentation", "Swagger/OpenAPI", "Tài liệu API"]
    ]

    create_table_with_header(doc, backend_headers, backend_data)

    doc.add_paragraph()

    doc.add_heading("3.2. Frontend Stack", level=2)

    frontend_headers = ["Component", "Technology", "Purpose"]
    frontend_data = [
        ["Framework", "Next.js 14", "Web application (App Router)"],
        ["Language", "TypeScript", "Type safety"],
        ["State Management", "TanStack Query", "Server state management"],
        ["UI Library", "Tailwind CSS", "Styling"],
        ["Component Library", "shadcn/ui", "Pre-built components"],
        ["Forms", "React Hook Form + Zod", "Form validation"]
    ]

    create_table_with_header(doc, frontend_headers, frontend_data)

    doc.add_paragraph()

    doc.add_heading("3.3. Cloud & Infrastructure", level=2)

    cloud_headers = ["Service", "Technology", "Purpose"]
    cloud_data = [
        ["Server", "AWS EC2", "Host tất cả services"],
        ["CDN", "Cloudflare", "Phân phối nội dung tĩnh"],
        ["Image Storage", "Cloudinary", "Lưu trữ ảnh bìa sách, avatar"],
        ["File Storage", "Cloudflare R2", "Lưu trữ file PDF/EPUB ebooks"],
        ["Email", "SendGrid/AWS SES", "Gửi email thông báo"]
    ]

    create_table_with_header(doc, cloud_headers, cloud_data)

    doc.add_paragraph()

    doc.add_heading("3.4. Thanh toán", level=2)

    payment_headers = ["Provider", "Purpose", "Status"]
    payment_data = [
        ["PayOS", "Thanh toán trực tuyến", "Tích hợp"],
        ["COD", "Thanh toán khi nhận hàng", "Tích hợp"]
    ]

    create_table_with_header(doc, payment_headers, payment_data)

    doc.add_page_break()

    # ==================== SECTION 4: DATABASE ====================
    doc.add_heading("4. THIẾT KẾ DATABASE", level=1)

    doc.add_heading("4.1. PostgreSQL Databases", level=2)

    db_intro = """Hệ thống sử dụng 5 PostgreSQL databases riêng biệt cho mỗi service, đảm bảo tính độc lập và dễ quản lý."""
    doc.add_paragraph(db_intro)

    postgres_headers = ["Database", "Service", "Tables"]
    postgres_data = [
        ["identity_db", "Identity Service", "users, sessions, refresh_tokens"],
        ["business_db", "Business Service", "businesses, stores, members"],
        ["commerce_db", "Commerce Service", "books, cart_items, orders, payments"],
        ["shipping_db", "Shipping Service", "shipments, delivery_logs, addresses"],
        ["promotion_db", "Promotion Service", "vouchers, promotions, discounts"]
    ]

    create_table_with_header(doc, postgres_headers, postgres_data)

    doc.add_paragraph()

    doc.add_heading("4.2. MongoDB Database", level=2)

    mongo_intro = """Community Service sử dụng MongoDB để lưu trữ các dữ liệu phi cấu trúc như forum posts, chat messages, và reviews."""
    doc.add_paragraph(mongo_intro)

    mongo_headers = ["Collection", "Mô tả"]
    mongo_data = [
        ["forums", "Diễn đàn thảo luận"],
        ["conversations", "Cuộc trò chuyện chat"],
        ["messages", "Tin nhắn chat"],
        ["reviews", "Đánh giá sách và cửa hàng"],
        ["notifications", "Thông báo in-app"],
        ["reports", "Báo cáo vi phạm"]
    ]

    create_table_with_header(doc, mongo_headers, mongo_data)

    doc.add_page_break()

    # ==================== SECTION 5: DEPLOYMENT ====================
    doc.add_heading("5. DEPLOYMENT", level=1)

    doc.add_heading("5.1. Kiến trúc Infrastructure trên AWS", level=2)

    deploy_intro = """Hệ thống được triển khai trên AWS EC2 với kiến trúc Microservices, sử dụng các service cloud cho database, cache, và storage."""
    doc.add_paragraph(deploy_intro)

    # Deployment diagram
    deploy_diagram = """
┌─────────────────────────────────────────────────────────────────┐
│                         Internet                                  │
└─────────────────────────────┬───────────────────────────────────┘
                              │
                              ▼
                    ┌─────────────────┐
                    │  Cloudflare     │
                    │  (CDN, SSL)     │
                    └────────┬────────┘
                             │
                             ▼
                    ┌─────────────────┐
                    │  AWS CloudFront │
                    │  (Static Assets) │
                    └────────┬────────┘
                             │
                             ▼
┌─────────────────────────────────────────────────────────────────┐
│                      AWS EC2 (Server)                             │
│                                                                  │
│  ┌─────────┐ ┌─────────┐ ┌─────────┐ ┌─────────┐ ┌─────────┐  │
│  │ Gateway │ │Identity │ │Business │ │Commerce │ │Shipping │  │
│  │  :3000  │ │ :3001  │ │ :3002  │ │ :3003  │ │ :3004  │  │
│  └─────────┘ └─────────┘ └─────────┘ └─────────┘ └─────────┘  │
│  ┌─────────┐ ┌─────────┐                                     │
│  │Community│ │Promo-   │                                     │
│  │ :3005   │ │tion:3007│                                     │
│  └─────────┘ └─────────┘                                     │
└─────────────────────────────────────────────────────────────────┘
                             │
        ┌─────────────────────┼─────────────────────┐
        │                     │                     │
        ▼                     ▼                     ▼
┌───────────────┐     ┌───────────────┐     ┌───────────────┐
│  PostgreSQL    │     │   MongoDB     │     │    Redis      │
│  (AWS RDS)     │     │  (MongoDB     │     │  (AWS        │
│                │     │   Atlas)       │     │   ElastiCache)│
│ • identity_db  │     │               │     │               │
│ • business_db │     │ • community_db│     │ • Sessions   │
│ • commerce_db │     │               │     │ • Cache      │
│ • shipping_db  │     │               │     │ • Rate Limit │
│ • promotion_db│     │               │     │               │
└───────────────┘     └───────────────┘     └───────────────┘
"""
    deploy_para = doc.add_paragraph()
    deploy_run = deploy_para.add_run(deploy_diagram)
    deploy_run.font.name = 'Courier New'
    deploy_run.font.size = Pt(8)

    doc.add_heading("5.2. External Cloud Services", level=2)

    external_headers = ["Service", "Purpose", "Files stored"]
    external_data = [
        ["Cloudinary", "Lưu trữ hình ảnh", "Book covers, avatars, store logos, forum images"],
        ["Cloudflare R2", "Lưu trữ file lớn", "PDF ebooks, EPUB files, large documents"],
        ["PayOS", "Thanh toán", "Payment gateway integration"]
    ]

    create_table_with_header(doc, external_headers, external_data)

    doc.add_paragraph()

    doc.add_heading("5.3. Security Architecture", level=2)

    security_headers = ["Layer", "Implementation"]
    security_data = [
        ["Authentication", "JWT Access Token (15 phút) + Refresh Token (7 ngày)"],
        ["Authorization", "RBAC (Role-Based Access Control)"],
        ["Session Management", "Redis cache với TTL"],
        ["Rate Limiting", "NestJS Throttler (100 requests/phút)"],
        ["Password Security", "Bcrypt với 12 rounds"],
        ["API Security", "HTTPS, CORS, Input validation (Zod)"],
        ["Database", "SSL connections"]
    ]

    create_table_with_header(doc, security_headers, security_data)

    doc.add_page_break()

    # ==================== SECTION 6: CONCLUSION ====================
    doc.add_heading("6. KẾT LUẬN", level=1)

    doc.add_heading("6.1. Tổng kết", level=2)

    conclusion_text = """Đề tài "Hệ thống thương mại điện tử hỗ trợ bán sách vật lý và sách điện tử" mang đến giải pháp toàn diện cho việc mua bán sách trực tuyến tại Việt Nam. Hệ thống được xây dựng với kiến trúc Microservices hiện đại, sử dụng các công nghệ tiên tiến như NestJS, Next.js 14, PostgreSQL, MongoDB, đảm bảo khả năng mở rộng và bảo trì dễ dàng.

Các điểm nổi bật của hệ thống:
• Kiến trúc Microservices với 6 services độc lập
• Hỗ trợ đồng thời sách vật lý và sách điện tử
• Marketplace đa cửa hàng
• Thanh toán tích hợp PayOS
• Cloud infrastructure với AWS EC2, Cloudflare, Cloudinary, R2"""
    doc.add_paragraph(conclusion_text)

    doc.add_heading("6.2. Hướng phát triển tương lai", level=2)

    future_items = [
        "Phát triển Mobile Application (Flutter)",
        "Tích hợp thêm cổng thanh toán (VNPay, MoMo)",
        "Triển khai AI cho gợi ý sách",
        "Hệ thống khuyến nghị cá nhân hóa",
        "Tính năng audiobook"
    ]

    for item in future_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    # References
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)

    refs = [
        "1. NestJS Documentation - https://docs.nestjs.com",
        "2. Next.js Documentation - https://nextjs.org/docs",
        "3. Prisma Documentation - https://pris.dev",
        "4. AWS EC2 Documentation - https://docs.aws.amazon.com/ec2",
        "5. Cloudflare Documentation - https://developers.cloudflare.com",
        "6. PayOS Documentation - https://docs.payos.vn"
    ]

    for ref in refs:
        doc.add_paragraph(ref)

    # Save document
    output_path = r"C:\Users\acer\Desktop\HUKI_EBOOK_BaoCao.docx"
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    main()
